"""Generation stage (Phase 2 scaffold): GMC-format CV as .docx.

Section list, order, and required fields come from `rules/cv_structure.json`
(HARD RULE 5) — so this refuses to run in production until that file is verified
line-by-line against the GMC CV guidance. Employment is reverse-chronological
back to primary medical qualification; gaps are listed explicitly.

HARD RULE 1: blank source fields stay blank in the CV and are surfaced in the
gap report — never filled in.
"""
from __future__ import annotations

from io import BytesIO

from app.models import Applicant, CareerEvent
from app.pipeline.aggregate import detect_gaps
from app.rules.loader import load_rules

_MISSING = ""  # blank, never inferred


def _fmt_date(d) -> str:
    return d.isoformat() if d else _MISSING


def build_cv_docx(applicant: Applicant, events: list[CareerEvent]) -> bytes:
    """Render a GMC-format CV to .docx bytes. Raises UnverifiedRulesError in
    production until cv_structure.json is verified."""
    from docx import Document as Docx  # lazy import

    rules = load_rules("cv_structure")["rules"]
    min_gap = int(rules.get("gaps", {}).get("min_gap_days_to_flag", 28))

    doc = Docx()
    doc.add_heading("Curriculum Vitae", level=0)

    for section in rules.get("sections", []):
        doc.add_heading(section.get("label", section["id"]), level=1)

        if section["id"] == "personal_details":
            doc.add_paragraph(f"Name: {applicant.full_name or _MISSING}")
            doc.add_paragraph(f"GMC reference: {applicant.gmc_reference_number or _MISSING}")

        elif section["id"] == "employment_history":
            dated = sorted(
                (e for e in events if not e.is_gap),
                key=lambda e: (e.start_date or _min_date(), e.end_date or _min_date()),
                reverse=(section.get("ordering", "reverse_chronological") == "reverse_chronological"),
            )
            for e in dated:
                line = (
                    f"{_fmt_date(e.start_date)} – {_fmt_date(e.end_date) or 'present'}: "
                    f"{e.post_title or _MISSING}"
                )
                p = doc.add_paragraph(line, style="List Bullet")
                details = ", ".join(
                    x for x in [e.grade, e.employer,
                                f"{e.wte_percent}% WTE" if e.wte_percent is not None else None,
                                e.employment_type.value if e.employment_type else None]
                    if x
                )
                if details:
                    p.add_run(f"\n    {details}")

        elif section["id"] == "gaps_in_employment":
            gaps = detect_gaps(events, min_gap_days=min_gap)
            if not gaps:
                doc.add_paragraph("No gaps in employment detected.")
            for g in gaps:
                doc.add_paragraph(
                    f"{g.start_date.isoformat()} – {g.end_date.isoformat()} "
                    f"({g.days} days) — reason: {_MISSING}",
                    style="List Bullet",
                )

        elif section["id"] == "qualifications":
            doc.add_paragraph(
                "[Qualifications extracted from certificates appear here — "
                "blank fields are surfaced in the gap report.]"
            )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _min_date():
    from datetime import date
    return date.min
